#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Cowork 沙盒網路存取安全評估報告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('外部域名開放可行性分析與建議')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.italic = True
subtitle_format.font.size = Pt(12)

# Date
meta = doc.add_paragraph(f'準備日期：{datetime.now().strftime("%Y年%m月%d日")} | 評估者：Ted (資安管理)')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_format = meta.runs[0]
meta_format.font.size = Pt(10)
meta_format.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph()

# Executive Summary
doc.add_heading('執行摘要', level=1)
doc.add_paragraph('本評估涵蓋8個外部域名的安全性和業務必要性，評分結果如下：')

summary_items = [
    ('建議開放（低風險）：5個域名 - GitHub、FinMind、Yahoo Finance', RGBColor(0, 176, 80)),
    ('有條件開放（中風險）：2個域名 - Yahoo 股票、經濟日報（需額外控制）', RGBColor(255, 192, 0)),
    ('不建議開放（高風險）：1個域名 - YouTube（高度可被利用的內容平台）', RGBColor(192, 0, 0))
]

for text, color in summary_items:
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = color

doc.add_paragraph()

# Assessment Table
doc.add_heading('詳細評估結果', level=1)

table = doc.add_table(rows=9, cols=6)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
headers = ['域名', '業務必要性', '風險等級', '認證模式', '推薦決策', '優先度']
for i, header in enumerate(headers):
    cell = header_cells[i]
    cell.text = header
    # Style header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '2E75B6')
    cell._element.get_or_add_tcPr().append(shading_elm)

# Data rows
data = [
    ('github.com', '高\n(CI/CD 核心)', '低', 'OAuth 2.0\nPAT', '開放 ✓', 'P0\n立即', RGBColor(0, 176, 80)),
    ('api.github.com', '高\n(GitHub API)', '低', 'OAuth 2.0\nPAT', '開放 ✓', 'P0\n立即', RGBColor(0, 176, 80)),
    ('raw.githubusercontent.com', '高\n(原始檔案)', '低', '無（公開）', '開放 ✓', 'P0\n立即', RGBColor(0, 176, 80)),
    ('finmindapi.site\nfinmindtrade.com', '高\n(選股資料)', '低', 'API 密鑰', '開放 ✓', 'P0\n立即', RGBColor(0, 176, 80)),
    ('query2.finance.yahoo.com', '高\n(粗篩過濾)', '低', '無（公開）', '開放 ✓', 'P0\n立即', RGBColor(0, 176, 80)),
    ('www.youtube.com', '低\n(用途模糊)', '高', '複雜\n(通用網站)', '不開放 ✗', '拒絕\n改用 API', RGBColor(192, 0, 0)),
    ('tw.stock.yahoo.com', '中\n(新聞資訊)', '中', '無（爬蟲）', '條件開放 ⚠', 'P1\n評估合規', RGBColor(255, 192, 0)),
    ('money.udn.com', '中\n(財經新聞)', '中', '無（爬蟲）', '條件開放 ⚠', 'P2\n評估合規', RGBColor(255, 192, 0)),
]

for row_idx, (domain, necessity, risk, auth, decision, priority, color) in enumerate(data, 1):
    row = table.rows[row_idx]
    cells = row.cells

    cells[0].text = domain
    cells[1].text = necessity
    cells[2].text = risk
    cells[3].text = auth
    cells[4].text = decision
    cells[5].text = priority

    # Color decision cell
    cells[4].paragraphs[0].runs[0].font.bold = True
    cells[4].paragraphs[0].runs[0].font.color.rgb = color

doc.add_paragraph()

# Detailed Analysis
doc.add_heading('各域名詳細分析', level=1)

# GitHub
doc.add_heading('1. github.com / api.github.com / raw.githubusercontent.com', level=2)
doc.add_paragraph('建議：強烈建議開放 ✓', style='List Bullet')
p = doc.add_paragraph('業務必要性：Moly 自動排程依賴 gh CLI 觸發 GitHub Actions，是核心 CI/CD 流程', style='List Bullet')
p.runs[0].bold = True
doc.add_paragraph('風險等級：低風險。GitHub 是業界標準的代碼管理平台，擁有企業級安全防護', style='List Bullet')
doc.add_paragraph('安全特性：', style='List Bullet')
doc.add_paragraph('- OAuth 2.0 認證，個人存取令牌需明確授權', style='List Number')
doc.add_paragraph('- raw.githubusercontent.com 用於檢索已合併程式碼，風險受控', style='List Number')
doc.add_paragraph('- 所有操作留有完整審計日誌', style='List Number')
doc.add_paragraph('建議控制：使用組織級別的 fine-grained Personal Access Token，限制範圍到特定儲存庫和操作', style='List Bullet')

# FinMind
doc.add_heading('2. finmindapi.site / finmindtrade.com', level=2)
doc.add_paragraph('建議：建議開放 ✓', style='List Bullet')
p = doc.add_paragraph('業務必要性：選股模型必需的台股資料來源，提供歷史價格、基本面數據', style='List Bullet')
p.runs[0].bold = True
doc.add_paragraph('風險等級：低風險（只讀 API）', style='List Bullet')
doc.add_paragraph('安全特性：', style='List Bullet')
doc.add_paragraph('- API 密鑰認證模式，不涉及登入會話', style='List Number')
doc.add_paragraph('- 台灣知名的開源財務資料平台，社群維護', style='List Number')
doc.add_paragraph('- 傳輸的是公開市場數據，無敏感性', style='List Number')
doc.add_paragraph('建議控制：使用環境變數管理 API 密鑰；限制 API 呼叫速率（避免濫用導致 IP 被封禁）', style='List Bullet')

# Yahoo Finance
doc.add_heading('3. query2.finance.yahoo.com（Yahoo Finance API）', level=2)
doc.add_paragraph('建議：建議開放 ✓', style='List Bullet')
p = doc.add_paragraph('業務必要性：選股粗篩階段的全球市場數據來源，包含股票代碼驗證和基本報價', style='List Bullet')
p.runs[0].bold = True
doc.add_paragraph('風險等級：低風險（只讀、無認證 API）', style='List Bullet')
doc.add_paragraph('安全特性：', style='List Bullet')
doc.add_paragraph('- 公開的財務數據，無身份驗證需求', style='List Number')
doc.add_paragraph('- Yahoo 是全球知名公司，基礎設施成熟', style='List Number')
doc.add_paragraph('- 傳輸的資訊完全是公開市場數據', style='List Number')
doc.add_paragraph('建議控制：實施速率限制防止濫用；監測異常的大量查詢；使用 User-Agent 識別應用程式', style='List Bullet')

# YouTube
doc.add_heading('4. www.youtube.com', level=2)
doc.add_paragraph('建議：不建議開放 ✗', style='List Bullet')
p = doc.add_paragraph('業務必要性：影片資訊讀取 - 用途不清楚，且與選股核心邏輯無直接關聯', style='List Bullet')
p.runs[0].bold = True
doc.add_paragraph('風險等級：高風險（通用內容平台，高度可被利用）', style='List Bullet')
doc.add_paragraph('安全隱憂：', style='List Bullet')
doc.add_paragraph('- 通用媒體平台，容易成為數據外洩、跟蹤、廣告追蹤的源頭', style='List Number')
doc.add_paragraph('- 可能加載追蹤像素、第三方指令碼，造成沙盒逃逸風險', style='List Number')
doc.add_paragraph('- 用途模糊，難以制定有效的存取控制和監測', style='List Number')
doc.add_paragraph('替代方案：若需要影片中繼資訊，建議透過 YouTube Data API（有明確的認證和配額限制）而非網站爬蟲', style='List Bullet')

# Yahoo Stock
doc.add_heading('5. tw.stock.yahoo.com（Yahoo 台股）', level=2)
doc.add_paragraph('建議：有條件開放 ⚠', style='List Bullet')
p = doc.add_paragraph('業務必要性：台股新聞和實時資訊，輔助選股決策', style='List Bullet')
p.runs[0].bold = True
doc.add_paragraph('風險等級：中風險（網站爬蟲不明確、可能違反服務條款）', style='List Bullet')
doc.add_paragraph('安全隱憂：', style='List Bullet')
doc.add_paragraph('- 網站內容爬蟲行為可能違反 ToS，面臨被封禁風險', style='List Number')
doc.add_paragraph('- 可能包含追蹤程式碼、廣告、第三方指令碼', style='List Number')
doc.add_paragraph('- 網站變動會影響爬蟲穩定性', style='List Number')
doc.add_paragraph('建議控制：優先評估 Yahoo Finance API 的台股支援；若必須爬蟲，使用 User-Agent 識別、設定適當延遲、監測被封禁風險；定期檢查選取器是否仍有效', style='List Bullet')

# UDN Money
doc.add_heading('6. money.udn.com（經濟日報財經新聞）', level=2)
doc.add_paragraph('建議：有條件開放 ⚠', style='List Bullet')
p = doc.add_paragraph('業務必要性：台灣本地財經新聞來源，提供市場情報和產業分析', style='List Bullet')
p.runs[0].bold = True
doc.add_paragraph('風險等級：中風險（新聞媒體網站，爬蟲治理模糊）', style='List Bullet')
doc.add_paragraph('安全隱憂：', style='List Bullet')
doc.add_paragraph('- 新聞網站爬蟲通常無明確 API，著作權風險', style='List Number')
doc.add_paragraph('- 可能包含廣告、分析追蹤碼、第三方指令碼', style='List Number')
doc.add_paragraph('- 網站頻繁更新，爬蟲維護成本高', style='List Number')
doc.add_paragraph('建議控制：評估使用 RSS Feed 或官方內容合作；若進行爬蟲，限制速率、遵守 robots.txt、監測封禁風險；建立合規管理，確保符合著作權法規', style='List Bullet')

doc.add_page_break()

# Security Recommendations
doc.add_heading('整體安全建議', level=1)

doc.add_heading('優先開放名單（Phase 1）', level=2)
doc.add_paragraph('github.com、api.github.com、raw.githubusercontent.com（CI/CD 核心依賴）', style='List Bullet')
doc.add_paragraph('finmindapi.site、finmindtrade.com（選股資料來源）', style='List Bullet')
doc.add_paragraph('query2.finance.yahoo.com（全球市場篩選）', style='List Bullet')

doc.add_heading('條件開放名單（Phase 2，需評估具體實作）', level=2)
doc.add_paragraph('tw.stock.yahoo.com（建議優先改用 API 或評估爬蟲合規）', style='List Bullet')
doc.add_paragraph('money.udn.com（建議評估 RSS / 合作方案，或限制爬蟲範圍）', style='List Bullet')

doc.add_heading('不建議開放名單', level=2)
doc.add_paragraph('www.youtube.com（高風險、用途模糊，改用 API 或使用替代資料來源）', style='List Bullet')

doc.add_heading('通用安全控制框架', level=2)

doc.add_heading('1. 認證與授權', level=3)
doc.add_paragraph('- API 密鑰儲存在 Secrets Manager 或環境變數，不硬編碼', style='List Number')
doc.add_paragraph('- GitHub 使用 fine-grained PAT，限制範圍和權限', style='List Number')
doc.add_paragraph('- 定期輪轉認證憑證', style='List Number')

doc.add_heading('2. 速率限制與使用配額', level=3)
doc.add_paragraph('- 在應用層實施速率限制，防止濫用', style='List Number')
doc.add_paragraph('- 監測 API 配額，設定告警', style='List Number')
doc.add_paragraph('- 記錄所有外部 API 呼叫', style='List Number')

doc.add_heading('3. 日誌與監測', level=3)
doc.add_paragraph('- 記錄所有外部網域存取（目標、時間、狀態碼、數據量）', style='List Number')
doc.add_paragraph('- 設定異常存取的告警（如超大數據傳輸、非正常時段存取）', style='List Number')
doc.add_paragraph('- 定期審計日誌，檢查未授權存取', style='List Number')

doc.add_heading('4. 網路隔離與 DNS', level=3)
doc.add_paragraph('- 設定防火牆白名單，明確限制允許的域名和 IP 範圍', style='List Number')
doc.add_paragraph('- 使用 DNS 過濾，防止 DNS 重新綁定攻擊', style='List Number')
doc.add_paragraph('- 考慮使用代理伺服器檢查出站流量', style='List Number')

doc.add_heading('5. 內容安全與資料驗證', level=3)
doc.add_paragraph('- 驗證響應內容的格式和簽名（如 JSON 結構、SSL 憑證驗證）', style='List Number')
doc.add_paragraph('- 在沙盒環境中隔離外部資料，防止注入攻擊', style='List Number')
doc.add_paragraph('- 定期更新相依套件，修補已知漏洞', style='List Number')

doc.add_heading('6. 定期審查', level=3)
doc.add_paragraph('- 每季審查已開放域名的使用情況和安全事件', style='List Number')
doc.add_paragraph('- 若有新增域名需求，應重新評估並獲得資安簽核', style='List Number')

doc.add_page_break()

# Implementation Roadmap
doc.add_heading('實施路線圖', level=1)

roadmap_table = doc.add_table(rows=4, cols=4)
roadmap_table.style = 'Light Grid Accent 1'

# Header
roadmap_header = roadmap_table.rows[0].cells
roadmap_headers = ['階段', '時程', '行動項目', '負責人']
for i, header in enumerate(roadmap_headers):
    cell = roadmap_header[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '2E75B6')
    cell._element.get_or_add_tcPr().append(shading_elm)

# Phase 1
phase1_cells = roadmap_table.rows[1].cells
phase1_cells[0].text = 'Phase 1\n立即執行'
phase1_cells[1].text = '第1週'
phase1_cells[2].text = '• 開放 GitHub 三域名\n• 開放 FinMind 二域名\n• 開放 Yahoo Finance\n• 部署基本監測'
phase1_cells[3].text = '基礎設施\n+ 資安'

# Phase 2
phase2_cells = roadmap_table.rows[2].cells
phase2_cells[0].text = 'Phase 2\n條件評估'
phase2_cells[1].text = '第2-3週'
phase2_cells[2].text = '• 評估 tw.stock.yahoo.com 爬蟲合規\n• 評估 money.udn.com 合規或替代方案\n• 建立爬蟲治理政策'
phase2_cells[3].text = '產品 + 資安\n+ 法務'

# Phase 3
phase3_cells = roadmap_table.rows[3].cells
phase3_cells[0].text = 'Phase 3\n監測與調整'
phase3_cells[1].text = '持續進行'
phase3_cells[2].text = '• 監測外部 API 使用\n• 季度安全審計\n• 評估新增域名需求'
phase3_cells[3].text = '資安 + 基礎設施'

doc.add_paragraph()

# Closing
closing = doc.add_paragraph()
closing.add_run('本評估報告提供資安角度的建議，最終決策應由基礎設施和產品團隊協商確定。如有任何疑問或需進一步評估，請聯繫資安管理團隊。').italic = True

# Save document
doc.save('/sessions/cool-dreamy-faraday/mnt/桌面/Cowork_Network_Access_Security_Assessment.docx')
print("✓ 評估報告已成功生成！")
print("📄 文件位置：C:\\Users\\wangj\\OneDrive\\桌面\\Cowork_Network_Access_Security_Assessment.docx")
